# ============================================================
# knowledge/file_reader.py
# ============================================================
# RESPONSABILIDAD:
# Capa inicial de lectura de archivos para NIA.
#
# Este módulo pertenece a la futura capa:
# - module_vision_archivos
# - RAG documental
# - lectura de fichas técnicas
# - lectura de manuales
# - lectura de catálogos de proveedor
# - lectura de documentos comerciales o técnicos
#
# IMPORTANTE:
# Este módulo NO recomienda productos.
# Este módulo NO consulta el catálogo.
# Este módulo NO responde al usuario.
# Este módulo NO se conecta todavía al orquestador.
#
# Solo hace:
# 1. Detectar tipo de archivo.
# 2. Leer contenido.
# 3. Limpiar texto.
# 4. Dividir en chunks.
# 5. Devolver una estructura estándar.
#
# Regla de Don Andrés:
# - Catálogo MongoDB = fuente de verdad para productos.
# - Archivos técnicos = soporte para explicar, validar o complementar.
# ============================================================

from __future__ import annotations

import csv
import json
import re
from pathlib import Path
from typing import Any, Dict, List, Optional


# ============================================================
# CONFIGURACIÓN
# ============================================================

SUPPORTED_EXTENSIONS = {
    ".txt",
    ".md",
    ".json",
    ".csv",
    ".pdf",
    ".docx",
    ".xlsx",
}

DEFAULT_CHUNK_SIZE = 1200
DEFAULT_CHUNK_OVERLAP = 150


# ============================================================
# UTILIDADES BÁSICAS
# ============================================================

def _safe_str(value: Any) -> str:
    """
    Convierte cualquier valor a texto seguro.
    """
    if value is None:
        return ""

    try:
        return str(value)
    except Exception:
        return ""


def clean_text(text: str) -> str:
    """
    Limpia texto extraído de archivos.

    Objetivo:
    - Quitar espacios excesivos.
    - Quitar saltos repetidos.
    - Mantener texto legible.
    """
    text = _safe_str(text)

    # Normaliza saltos de línea.
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # Quita espacios repetidos por línea.
    lines = []
    for line in text.split("\n"):
        cleaned_line = re.sub(r"[ \t]+", " ", line).strip()
        if cleaned_line:
            lines.append(cleaned_line)

    # Une dejando saltos simples.
    cleaned = "\n".join(lines)

    # Evita demasiados saltos consecutivos.
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)

    return cleaned.strip()


def detect_file_type(file_path: str | Path) -> Dict[str, Any]:
    """
    Detecta el tipo de archivo por extensión.
    """
    path = Path(file_path)
    extension = path.suffix.lower()

    return {
        "file_name": path.name,
        "file_path": str(path),
        "extension": extension,
        "supported": extension in SUPPORTED_EXTENSIONS,
    }


def _build_error_response(
    file_path: str | Path,
    error_message: str,
    file_type: Optional[str] = None,
) -> Dict[str, Any]:
    """
    Construye respuesta estándar de error.
    """
    path = Path(file_path)

    return {
        "ok": False,
        "file_path": str(path),
        "file_name": path.name,
        "file_type": file_type or path.suffix.lower(),
        "text": "",
        "chunks": [],
        "metadata": {
            "chars": 0,
            "chunk_count": 0,
        },
        "errors": [error_message],
    }


# ============================================================
# LECTORES POR TIPO DE ARCHIVO
# ============================================================

def read_text_file(file_path: str | Path) -> str:
    """
    Lee archivos TXT o MD.
    """
    path = Path(file_path)

    # Primero intenta UTF-8.
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        # Fallback común en Windows / archivos antiguos.
        return path.read_text(encoding="latin-1")


def read_json_file(file_path: str | Path) -> str:
    """
    Lee un JSON y lo convierte a texto legible.

    Esto sirve tanto para:
    - módulos NIA OS
    - reglas internas
    - configuraciones
    - documentos estructurados
    """
    path = Path(file_path)

    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except UnicodeDecodeError:
        data = json.loads(path.read_text(encoding="latin-1"))

    return json.dumps(
        data,
        ensure_ascii=False,
        indent=2,
    )


def read_csv_file(file_path: str | Path) -> str:
    """
    Lee un CSV y lo convierte a texto plano tabular.
    """
    path = Path(file_path)

    rows: List[str] = []

    try:
        file_handle = path.open("r", encoding="utf-8", newline="")
    except UnicodeDecodeError:
        file_handle = path.open("r", encoding="latin-1", newline="")

    with file_handle as csv_file:
        reader = csv.reader(csv_file)

        for row in reader:
            rows.append(" | ".join(_safe_str(cell) for cell in row))

    return "\n".join(rows)


def read_pdf_file(file_path: str | Path) -> str:
    """
    Lee texto de un PDF.

    Requiere:
    pip install pypdf

    Nota:
    Esta función extrae texto. No hace OCR.
    Si el PDF es escaneado como imagen, puede devolver poco o nada.
    """
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ImportError(
            "No está instalada la librería pypdf. Instálala con: pip install pypdf"
        ) from exc

    path = Path(file_path)
    reader = PdfReader(str(path))

    pages_text = []

    for page_index, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""

        if page_text.strip():
            pages_text.append(
                f"[PÁGINA {page_index}]\n{page_text}"
            )

    return "\n\n".join(pages_text)


def read_docx_file(file_path: str | Path) -> str:
    """
    Lee texto de un archivo DOCX.

    Requiere:
    pip install python-docx
    """
    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError(
            "No está instalada la librería python-docx. Instálala con: pip install python-docx"
        ) from exc

    path = Path(file_path)
    document = Document(str(path))

    parts = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if text:
            parts.append(text)

    # También intenta leer tablas.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            cells = [cell for cell in cells if cell]

            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def read_xlsx_file(file_path: str | Path) -> str:
    """
    Lee texto de un archivo XLSX.

    Requiere:
    pip install openpyxl

    Nota:
    Esta función convierte filas a texto.
    podemos crear un lector especializado para catálogos.
    """
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise ImportError(
            "No está instalada la librería openpyxl. Instálala con: pip install openpyxl"
        ) from exc

    path = Path(file_path)

    workbook = load_workbook(
        filename=str(path),
        read_only=True,
        data_only=True,
    )

    parts = []

    for sheet in workbook.worksheets:
        parts.append(f"[HOJA: {sheet.title}]")

        for row in sheet.iter_rows(values_only=True):
            values = [_safe_str(value).strip() for value in row]
            values = [value for value in values if value]

            if values:
                parts.append(" | ".join(values))

    return "\n".join(parts)


# ============================================================
# CHUNKING
# ============================================================

def split_into_chunks(
    text: str,
    chunk_size: int = DEFAULT_CHUNK_SIZE,
    overlap: int = DEFAULT_CHUNK_OVERLAP,
    source: Optional[str] = None,
) -> List[Dict[str, Any]]:
    """
    Divide texto en chunks para futura indexación/retrieval.

    Estrategia simple:
    - Divide por tamaño de caracteres.
    - Mantiene un pequeño solapamiento entre chunks.
    - No usa embeddings todavía.
    """
    text = clean_text(text)

    if not text:
        return []

    if chunk_size <= 0:
        chunk_size = DEFAULT_CHUNK_SIZE

    if overlap < 0:
        overlap = 0

    if overlap >= chunk_size:
        overlap = max(0, chunk_size // 5)

    chunks = []
    start = 0
    chunk_number = 1

    while start < len(text):
        end = start + chunk_size
        chunk_text = text[start:end].strip()

        if chunk_text:
            chunks.append({
                "chunk_id": f"chunk_{chunk_number:04d}",
                "text": chunk_text,
                "source": source or "",
                "start_char": start,
                "end_char": min(end, len(text)),
                "chars": len(chunk_text),
            })

            chunk_number += 1

        if end >= len(text):
            break

        start = end - overlap

    return chunks


# ============================================================
# FUNCIÓN PRINCIPAL
# ============================================================

def read_knowledge_file(
    file_path: str | Path,
    chunk_size: int = DEFAULT_CHUNK_SIZE,
    overlap: int = DEFAULT_CHUNK_OVERLAP,
) -> Dict[str, Any]:
    """
    Lee un archivo de conocimiento y devuelve estructura estándar.

    Esta es la función principal que usaremos más adelante desde:
    - scripts de ingesta
    - panel admin
    - carga de documentos
    - RAG documental
    """
    path = Path(file_path)

    if not path.exists():
        return _build_error_response(
            file_path=path,
            error_message=f"El archivo no existe: {path}",
        )

    if not path.is_file():
        return _build_error_response(
            file_path=path,
            error_message=f"La ruta no corresponde a un archivo: {path}",
        )

    file_info = detect_file_type(path)
    extension = file_info["extension"]

    if not file_info["supported"]:
        return _build_error_response(
            file_path=path,
            file_type=extension,
            error_message=f"Tipo de archivo no soportado todavía: {extension}",
        )

    try:
        if extension in [".txt", ".md"]:
            raw_text = read_text_file(path)

        elif extension == ".json":
            raw_text = read_json_file(path)

        elif extension == ".csv":
            raw_text = read_csv_file(path)

        elif extension == ".pdf":
            raw_text = read_pdf_file(path)

        elif extension == ".docx":
            raw_text = read_docx_file(path)

        elif extension == ".xlsx":
            raw_text = read_xlsx_file(path)

        else:
            return _build_error_response(
                file_path=path,
                file_type=extension,
                error_message=f"No hay lector implementado para: {extension}",
            )

        cleaned_text = clean_text(raw_text)

        chunks = split_into_chunks(
            text=cleaned_text,
            chunk_size=chunk_size,
            overlap=overlap,
            source=path.name,
        )

        return {
            "ok": True,
            "file_path": str(path),
            "file_name": path.name,
            "file_type": extension,
            "text": cleaned_text,
            "chunks": chunks,
            "metadata": {
                "chars": len(cleaned_text),
                "chunk_count": len(chunks),
                "chunk_size": chunk_size,
                "overlap": overlap,
            },
            "errors": [],
        }

    except Exception as error:
        return _build_error_response(
            file_path=path,
            file_type=extension,
            error_message=str(error),
        )


# ============================================================
# LECTURA DE CARPETA
# ============================================================

def read_knowledge_folder(
    folder_path: str | Path,
    recursive: bool = True,
    chunk_size: int = DEFAULT_CHUNK_SIZE,
    overlap: int = DEFAULT_CHUNK_OVERLAP,
) -> Dict[str, Any]:
    """
    Lee todos los archivos soportados dentro de una carpeta.

    Esto será útil para procesar una carpeta completa de conocimiento.
    """
    folder = Path(folder_path)

    if not folder.exists():
        return {
            "ok": False,
            "folder_path": str(folder),
            "files": [],
            "summary": {
                "total_files": 0,
                "ok_files": 0,
                "error_files": 0,
                "total_chunks": 0,
            },
            "errors": [f"La carpeta no existe: {folder}"],
        }

    if not folder.is_dir():
        return {
            "ok": False,
            "folder_path": str(folder),
            "files": [],
            "summary": {
                "total_files": 0,
                "ok_files": 0,
                "error_files": 0,
                "total_chunks": 0,
            },
            "errors": [f"La ruta no corresponde a una carpeta: {folder}"],
        }

    pattern = "**/*" if recursive else "*"
    file_paths = [
        item for item in folder.glob(pattern)
        if item.is_file() and item.suffix.lower() in SUPPORTED_EXTENSIONS
    ]

    results = []

    for file_path in file_paths:
        result = read_knowledge_file(
            file_path=file_path,
            chunk_size=chunk_size,
            overlap=overlap,
        )

        results.append(result)

    ok_files = sum(1 for item in results if item.get("ok"))
    error_files = len(results) - ok_files
    total_chunks = sum(
        len(item.get("chunks", []))
        for item in results
        if isinstance(item, dict)
    )

    return {
        "ok": error_files == 0,
        "folder_path": str(folder),
        "files": results,
        "summary": {
            "total_files": len(results),
            "ok_files": ok_files,
            "error_files": error_files,
            "total_chunks": total_chunks,
        },
        "errors": [
            error
            for item in results
            for error in item.get("errors", [])
        ],
    }